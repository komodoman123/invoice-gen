"""
Invoice Generator & Email Sender
=================================
Reads a CSV, fills a .docx invoice template, converts to PDF, optionally sends via Gmail.

Setup:
    1.  pip install -r requirements.txt
    2.  Copy .env.example → .env  and fill in Gmail credentials
    3.  Edit the CONFIG block below (Sheet URL, column names)
    4.  python3 -m streamlit run app.py
"""

import os
import smtplib
from email import encoders
from email.mime.base import MIMEBase
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText

import tempfile

import pythoncom

import pandas as pd
import streamlit as st
from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx2pdf import convert as docx2pdf_convert
from dotenv import load_dotenv
from num2words import num2words

load_dotenv()


# ============================================================
# CONFIG  ← edit this block
# ============================================================

DOCX_TEMPLATE = "template.docx"  # your template file
EMAIL_COLUMN  = "E-mail"                          # CSV column with recipient addresses

# Marker name (docx)  →  Column name (CSV)
# Adjust the RIGHT side to match your actual sheet headers.
COLUMNS = {
    "Client":        "Client",
    "PIC":           "PIC",
    "Contact":       "Contact",
    "E-mail":        "E-mail",
    "Address":       "Address",
    "No. Invoice":   "No. Invoice",
    "No. Quotation": "No. Quotation",
    "Invoice Date":  "Invoice Date",
    "Due Date":      "Due Date",
}
for _i in range(1, 7):
    COLUMNS[f"Service {_i}"] = f"Service {_i}"
    COLUMNS[f"Qty {_i}"]     = f"Qty {_i}"
    COLUMNS[f"Price {_i}"]   = f"Price {_i}"

# Gmail  ─  reads from .env (see .env.example)
try:                                            # Streamlit Cloud
    GMAIL_SENDER       = st.secrets["GMAIL_SENDER"]
    GMAIL_APP_PASSWORD = st.secrets["GMAIL_APP_PASSWORD"]
except Exception:                               # local .env
    GMAIL_SENDER       = os.getenv("GMAIL_SENDER", "")
    GMAIL_APP_PASSWORD = os.getenv("GMAIL_APP_PASSWORD", "")

EMAIL_SUBJECT = "Invoice"
EMAIL_BODY    = "Hi,\n\nPlease find your invoice attached.\n\nBest regards"


# ============================================================
# helpers
# ============================================================

def fmt_idr(value) -> str:
    """1500000  →  '1.500.000'"""
    try:
        return f"{int(round(float(value))):,}".replace(",", ".")
    except (ValueError, TypeError):
        return str(value)


def to_terbilang(total: float) -> str:
    """Number  →  Indonesian words + ' rupiah'."""
    try:
        words = num2words(int(round(total)), lang="id")
        return words.capitalize() + " rupiah"
    except Exception:
        return fmt_idr(total)


def safe(row: dict, key: str) -> str:
    """Lookup that turns NaN → empty string."""
    v = row.get(key, "")
    return "" if pd.isna(v) else str(v).strip()


# ============================================================
# invoice generation  (docx → PDF)
# ============================================================

def _replace_in_paragraph(paragraph, replacements: dict):
    """Replace placeholders in one paragraph, handling text split across runs."""
    full = paragraph.text
    new  = full
    for old, val in replacements.items():
        new = new.replace(old, val)
    if new != full and paragraph.runs:
        paragraph.runs[0].text = new
        for run in paragraph.runs[1:]:
            run.text = ""


def generate_pdf(row: dict) -> bytes:
    """Fill the .docx template with *row*, convert to PDF, return bytes.

    Note: docx2pdf requires Microsoft Word installed (Windows only).
    """
    doc = Document(DOCX_TEMPLATE)

    # ── calculated columns ──────────────────────────────
    subtotals, total = {}, 0.0
    for i in range(1, 7):
        if not safe(row, COLUMNS.get(f"Service {i}", "")):
            subtotals[i] = 0.0
            continue
        try:
            sub = (float(safe(row, COLUMNS.get(f"Qty {i}",   "")) or 0)
                 * float(safe(row, COLUMNS.get(f"Price {i}", "")) or 0))
        except ValueError:
            sub = 0.0
        subtotals[i] = sub
        total       += sub

    # ── build placeholder → value map ────────────────────
    replacements = {}

    for field in ["Client", "PIC", "Contact", "E-mail", "Address",
                  "No. Invoice", "No. Quotation", "Invoice Date", "Due Date"]:
        replacements["{{" + field + "}}"] = safe(row, COLUMNS.get(field, field))

    for i in range(1, 7):
        svc = safe(row, COLUMNS.get(f"Service {i}", ""))
        replacements[f"{{{{Service {i}}}}}"] = svc
        if svc:
            raw_qty = safe(row, COLUMNS.get(f"Qty {i}", ""))
            try:
                replacements[f"{{{{Qty {i}}}}}"] = str(int(float(raw_qty))) if raw_qty else ""
            except (ValueError, TypeError):
                replacements[f"{{{{Qty {i}}}}}"] = raw_qty
            replacements[f"{{{{Price {i}}}}}"]    = fmt_idr(safe(row, COLUMNS.get(f"Price {i}", "")))
            replacements[f"{{{{Subtotal {i}}}}}"] = fmt_idr(subtotals[i])
        else:
            replacements[f"{{{{Qty {i}}}}}"]      = ""
            replacements[f"{{{{Price {i}}}}}"]    = ""
            replacements[f"{{{{Subtotal {i}}}}}"] = ""

    replacements["{{Total}}"]     = fmt_idr(total)
    replacements["{{Terbilang}}"] = to_terbilang(total)

    # ── apply replacements to body, tables, and footers ──
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, replacements)
    for table in doc.tables:
        for tbl_row in table.rows:
            for cell in tbl_row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, replacements)
    for rel in doc.part.rels.values():
        if "footer" in rel.reltype or "header" in rel.reltype:
            for p_elem in rel.target_part.element.iter(qn("w:p")):
                _replace_in_paragraph(Paragraph(p_elem, rel.target_part), replacements)

    # ── save .docx, convert to PDF, return bytes ─────────
    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, "invoice.docx")
        pdf_path  = os.path.join(tmpdir, "invoice.pdf")
        doc.save(docx_path)
        pythoncom.CoInitialize()
        try:
            docx2pdf_convert(docx_path, pdf_path)
        finally:
            pythoncom.CoUninitialize()
        with open(pdf_path, "rb") as f:
            return f.read()


# ============================================================
# email
# ============================================================

def send_email(recipient: str, pdf_bytes: bytes, filename: str):
    """Send *pdf_bytes* as an attachment via Gmail SMTP."""
    msg = MIMEMultipart()
    msg["From"]    = GMAIL_SENDER
    msg["To"]      = recipient
    msg["Subject"] = EMAIL_SUBJECT
    msg.attach(MIMEText(EMAIL_BODY, "plain"))

    part = MIMEBase("application", "octet-stream")
    part.set_payload(pdf_bytes)
    encoders.encode_base64(part)
    part.add_header("Content-Disposition", f'attachment; filename="{filename}.pdf"')
    msg.attach(part)

    with smtplib.SMTP("smtp.gmail.com", 587, timeout=10) as srv:
        srv.starttls()
        srv.login(GMAIL_SENDER, GMAIL_APP_PASSWORD)
        srv.send_message(msg)


# ============================================================
# Streamlit UI
# ============================================================

st.set_page_config(page_title="Invoice Generator", layout="wide")
st.title("Invoice Generator & Sender")

# ── guards ────────────────────────────────────────────────
if not os.path.exists(DOCX_TEMPLATE):
    st.error(f"Template not found: `{DOCX_TEMPLATE}`")
    st.stop()

if not GMAIL_SENDER or not GMAIL_APP_PASSWORD:
    st.warning(
        "Gmail not configured — sending is disabled.  "
        "Add credentials to `.env` (see `.env.example`)."
    )

# ── test / calibration ────────────────────────────────────
with st.expander("Test mode — check layout before using real data"):
    if st.button("Generate test PDF"):
        _test = {
            "Client": "PT Test Company",       "PIC":     "John Doe",
            "Contact": "+62 812 3456 7890",    "E-mail":  "test@example.com",
            "Address": "Jalan Test 123, Bandung, Jawa Barat",
            "No. Invoice": "INV-2025-001",     "No. Quotation": "QUO-2025-001",
            "Invoice Date": "01 Januari 2025", "Due Date": "31 Januari 2025",
        }
        for _j in range(1, 4):                 # 3 sample services
            _test[f"Service {_j}"] = f"Layanan contoh {_j}"
            _test[f"Qty {_j}"]     = str(_j * 2)
            _test[f"Price {_j}"]   = str(500_000 * _j)
        for _j in range(4, 7):                 # empty rows
            _test[f"Service {_j}"] = _test[f"Qty {_j}"] = _test[f"Price {_j}"] = ""

        st.download_button(
            "Download test PDF",
            data=generate_pdf(_test),
            file_name="invoice_test.pdf",
            mime="application/pdf",
        )

# ── load CSV ──────────────────────────────────────────────
uploaded = st.file_uploader("Upload invoice CSV", type="csv")
if uploaded:
    file_id = (uploaded.name, uploaded.size)
    if st.session_state.get("_csv_id") != file_id:
        st.session_state["df"]     = pd.read_csv(uploaded)
        st.session_state["_csv_id"] = file_id

if "df" not in st.session_state:
    st.info("Upload a CSV file above to get started.")
    st.stop()

df = st.session_state["df"]
st.dataframe(df, use_container_width=True)

# ── row selector ──────────────────────────────────────────
_client_col = COLUMNS.get("Client", "Client")
selected    = st.multiselect(
    "Select invoices to generate",
    options=df.index.tolist(),
    format_func=lambda i: (
        f"{safe(df.iloc[i].to_dict(), _client_col)}  |  "
        f"{safe(df.iloc[i].to_dict(), EMAIL_COLUMN)}"
    ),
)
if not selected:
    st.stop()

st.dataframe(df.iloc[selected], use_container_width=True)

# clear stale downloads if the selection changed
if set(st.session_state.get("generated", {}).keys()) != set(selected):
    st.session_state.pop("generated", None)

# ── action buttons ────────────────────────────────────────
c1, c2 = st.columns(2)

# — generate & download ──
if c1.button("Generate PDFs", use_container_width=True):
    with st.spinner("Generating PDFs..."):
        gen = {}
        for idx in selected:
            row  = df.iloc[idx].to_dict()
            name = (safe(row, _client_col) or f"invoice_{idx}").replace(" ", "_")
            gen[idx] = (name, generate_pdf(row))
        st.session_state["generated"] = gen
    st.success(f"{len(gen)} PDF(s) ready.")

if st.session_state.get("generated"):
    for idx, (name, data) in st.session_state["generated"].items():
        st.download_button(
            f"Download {name}.pdf", data=data,
            file_name=f"{name}.pdf", mime="application/pdf",
            key=f"dl_{idx}",
        )

# — generate & send ──
if c2.button("Generate & Send Emails", use_container_width=True):
    if not GMAIL_SENDER or not GMAIL_APP_PASSWORD:
        st.error("Gmail not configured — check `.env`.")
    else:
        with st.spinner("Sending emails..."):
            results = []
            for idx in selected:
                row       = df.iloc[idx].to_dict()
                recipient = safe(row, EMAIL_COLUMN)
                name      = (safe(row, _client_col) or f"invoice_{idx}").replace(" ", "_")
                if not recipient:
                    results.append(("error", f"Row {idx} ({name}): no email address"))
                    continue
                try:
                    send_email(recipient, generate_pdf(row), name)
                    results.append(("success", recipient))
                except Exception as exc:
                    results.append(("error", f"{recipient}: {exc}"))
        for kind, msg in results:
            (st.success if kind == "success" else st.error)(f"Sent to {msg}" if kind == "success" else f"Failed — {msg}")
